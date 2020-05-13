from PIL import Image, ImageDraw, ImageFont
from docx.shared import Inches
import docx
import requests


# Part 1:- This sections takes the taco image, resizes it, writes text over it, and saves it for part 2.

image = Image.open('Taco_image.jpg')
image.thumbnail((800, 800))                                                     # Resize image to 800 pixels tall
image.save('small_taco.jpg')                                                    # Save resized image

img_draw = ImageDraw.Draw(image)
font = ImageFont.truetype('DejaVuSans.ttf', 37)                                 # font specs
img_draw.text([180, 30], 'Random Taco Cookbook', fill='black', font=font)       # write on image
image.save('small_taco_text.jpg')                                               # Save resized image with text


# Part 2:- This sections gets the json data from the taco api and prints to Microsoft word

# ------------------------ First Taco Recipe ----------------------------

url_n1 = 'https://taco-1150.herokuapp.com/random/?full_taco=true'               # save url of taco api
taco_a = requests.get(url_n1).json()

# Store the values for the keys (names) and (recipe) for the 5 sections

# Save values in name for 5 cat
seasoning_n = taco_a['seasoning']['name']
condiment_n = taco_a['condiment']['name']
mixin_n = taco_a['mixin']['name']
base_layer_n = taco_a['base_layer']['name']
shell_n = taco_a['shell']['name']

# Save values in recipe for 5 cat
seasoning_r = taco_a['seasoning']['recipe']
condiment_r = taco_a['condiment']['recipe']
mixin_r = taco_a['mixin']['recipe']
base_layer_r = taco_a['base_layer']['recipe']
shell_r = taco_a['shell']['recipe']

# This section handles writing the json data and image to the Word document
# and is repeated two more times to get a total of 3 taco recipes.

# Document initializations and copy image to word doc.
document = docx.Document()
document.add_heading('Random Taco Cookbook', level=0)                       # Title for document

document.add_picture('small_taco_text.jpg', width=Inches(6.0))              # Resize the image in Part 1 to fit Word

document.add_heading('Credits', level=1)                                    # Credits as a header
document.add_paragraph('•   Taco image: Photo by Ryan Concepcion on Unsplash')
document.add_paragraph('•   Tacos from: https://taco-1150.herokuapp.com/random/?full_taco=true')
document.add_paragraph('•   Code by: Abdurashid Sharmarke')

# This subsection writes the json data (5 categories w/ names and recipes) to Word.
document.add_heading(seasoning_n, level=1)                                  # Write names as headers
document.add_paragraph(seasoning_r)                                         # Write recipes as standard text

document.add_heading(condiment_n, level=1)
document.add_paragraph(condiment_r)

document.add_heading(mixin_n, level=1)
document.add_paragraph(mixin_r)

document.add_heading(base_layer_n, level=1)
document.add_paragraph(base_layer_r)

document.add_heading(shell_n, level=1)
document.add_paragraph(shell_r)

document.add_page_break()                                                     # Page break

# ------------------------ Second Taco Recipe ----------------------------
# Save json data in variable
url_n2 = 'https://taco-1150.herokuapp.com/random/?full_taco=true'
taco_b = requests.get(url_n2).json()

# Save values in name for 5 cat
seasoning_n = taco_b['seasoning']['name']
condiment_n = taco_b['condiment']['name']
mixin_n = taco_b['mixin']['name']
base_layer_n = taco_b['base_layer']['name']
shell_n = taco_b['shell']['name']

# Save values in recipe for 5 cat
seasoning_r = taco_b['seasoning']['recipe']
condiment_r = taco_b['condiment']['recipe']
mixin_r = taco_b['mixin']['recipe']
base_layer_r = taco_b['base_layer']['recipe']
shell_r = taco_b['shell']['recipe']

# This subsection writes the json data (5 categories w/ names and recipes) to Word.
document.add_heading(seasoning_n, level=1)                                          # Write names as headers
document.add_paragraph(seasoning_r)                                                 # Write recipes as standard text

document.add_heading(condiment_n, level=1)
document.add_paragraph(condiment_r)

document.add_heading(mixin_n, level=1)
document.add_paragraph(mixin_r)

document.add_heading(base_layer_n, level=1)
document.add_paragraph(base_layer_r)

document.add_heading(shell_n, level=1)
document.add_paragraph(shell_r)

document.add_page_break()


# ------------------------ Third Taco Recipe ----------------------------
# Save json data in variable
url_n3 = 'https://taco-1150.herokuapp.com/random/?full_taco=true'
taco_c = requests.get(url_n3).json()

# Save values in name for 5 cat
seasoning_n = taco_c['seasoning']['name']
condiment_n = taco_c['condiment']['name']
mixin_n = taco_c['mixin']['name']
base_layer_n = taco_c['base_layer']['name']
shell_n = taco_c['shell']['name']

# Save values in recipe for 5 cat
seasoning_r = taco_c['seasoning']['recipe']
condiment_r = taco_c['condiment']['recipe']
mixin_r = taco_c['mixin']['recipe']
base_layer_r = taco_c['base_layer']['recipe']
shell_r = taco_c['shell']['recipe']

# This subsection writes the json data (5 categories w/ names and recipes) to Word.
document.add_heading(seasoning_n, level=1)                                     # Write names as headers
document.add_paragraph(seasoning_r)                                            # Write recipes as standard text

document.add_heading(condiment_n, level=1)
document.add_paragraph(condiment_r)

document.add_heading(mixin_n, level=1)
document.add_paragraph(mixin_r)

document.add_heading(base_layer_n, level=1)
document.add_paragraph(base_layer_r)

document.add_heading(shell_n, level=1)
document.add_paragraph(shell_r)


document.save('Final Taco.docx')                                                # Save document as Final Taco.
